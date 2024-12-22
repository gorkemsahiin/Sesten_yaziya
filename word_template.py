from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_legal_document(text, title="Hukuki Metin"):
    doc = Document()
    
    # Başlık stili
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tarih ekle
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(datetime.datetime.now().strftime("%d/%m/%Y"))
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Ana metin
    doc.add_paragraph(text)
    
    return doc 