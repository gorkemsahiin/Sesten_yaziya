import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
import os
from datetime import datetime
from voice_recorder import VoiceRecorder
from PIL import Image, ImageTk
from docx import Document
import threading
import time

class VoiceRecordingApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Hukuki Ses Kayıt ve Dönüştürme Uygulaması")
        self.root.geometry("1000x700")
        self.root.configure(bg='#f0f0f0')

        # Çıktı dizini
        self.output_dir = "converted_text"
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

        # Dosya yolu
        self.transcript_path = os.path.join(self.output_dir, "transcription.txt")

        # Ana çerçeve
        self.main_frame = ttk.Frame(self.root, padding="10")
        self.main_frame.pack(fill=tk.BOTH, expand=True)

        # Stil oluştur
        self.style = ttk.Style()
        self.style.configure('Custom.TButton', 
                           padding=10, 
                           font=('Helvetica', 10, 'bold'))

        # Üst bilgi çerçevesi
        self.header_frame = ttk.Frame(self.main_frame)
        self.header_frame.pack(fill=tk.X, pady=10)

        self.title_label = ttk.Label(
            self.header_frame,
            text="Hukuki Metin Dönüştürücü",
            font=('Helvetica', 16, 'bold')
        )
        self.title_label.pack()

        # Kontrol butonları çerçevesi
        self.control_frame = ttk.Frame(self.main_frame)
        self.control_frame.pack(fill=tk.X, pady=10)

        # Kayıt butonları
        self.record_button = ttk.Button(
            self.control_frame,
            text="🎤 Kayda Başla",
            style='Custom.TButton',
            command=self.start_recording
        )
        self.record_button.pack(side=tk.LEFT, padx=5)

        self.stop_button = ttk.Button(
            self.control_frame,
            text="⏹ Kaydı Durdur",
            style='Custom.TButton',
            command=self.stop_recording,
            state='disabled'
        )
        self.stop_button.pack(side=tk.LEFT, padx=5)

        # Dosya işlemleri çerçevesi
        self.file_frame = ttk.LabelFrame(self.main_frame, text="Dosya İşlemleri", padding="10")
        self.file_frame.pack(fill=tk.X, pady=10)

        # Dosya adı girişi
        self.filename_label = ttk.Label(self.file_frame, text="Dosya Adı:")
        self.filename_label.pack(side=tk.LEFT, padx=5)

        self.filename_entry = ttk.Entry(self.file_frame, width=30)
        self.filename_entry.pack(side=tk.LEFT, padx=5)
        self.filename_entry.insert(0, f"hukuki_metin_{datetime.now().strftime('%Y%m%d')}")

        # Word'e aktar butonu
        self.word_button = ttk.Button(
            self.file_frame,
            text="📄 Word'e Aktar",
            style='Custom.TButton',
            command=self.save_to_word
        )
        self.word_button.pack(side=tk.LEFT, padx=5)

        # Metin alanı çerçevesi
        self.text_frame = ttk.LabelFrame(self.main_frame, text="Dönüştürülen Metin", padding="10")
        self.text_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        # Metin alanı ve kaydırma çubuğu
        self.text_scroll = ttk.Scrollbar(self.text_frame)
        self.text_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.text_area = tk.Text(
            self.text_frame,
            wrap=tk.WORD,
            font=('Arial', 11),
            yscrollcommand=self.text_scroll.set
        )
        self.text_area.pack(fill=tk.BOTH, expand=True)
        self.text_scroll.config(command=self.text_area.yview)

        # Durum çubuğu
        self.status_frame = ttk.Frame(self.main_frame)
        self.status_frame.pack(fill=tk.X, pady=5)

        self.status_label = ttk.Label(
            self.status_frame,
            text="Hazır",
            font=('Helvetica', 9)
        )
        self.status_label.pack(side=tk.LEFT)

        # Voice Recorder örneği ve kayıt durumu
        self.recorder = VoiceRecorder()
        self.is_recording = False

    def start_recording(self):
        if not self.is_recording:
            self.is_recording = True
            self.record_button.configure(state='disabled')
            self.stop_button.configure(state='normal')
            self.status_label.configure(text="Kayıt yapılıyor...")
            self.text_area.delete(1.0, tk.END)
            
            # Kayıt işlemini başlat
            self.recorder.start_recording()

    def stop_recording(self):
        """Kaydı durdurma ve transkripsiyon"""
        try:
            self.is_recording = False
            self.status_label.configure(text="Kayıt durduruluyor ve metin dönüştürülüyor...")
            self.stop_button.configure(state='disabled')
            
            # Kaydı durdur
            self.recorder.stop_recording()
            
            # Transkripsiyon dosyasını kontrol et ve oku
            if os.path.exists(self.transcript_path):
                with open(self.transcript_path, 'r', encoding='utf-8') as file:
                    transcribed_text = file.read().strip()
                
                if transcribed_text:
                    # Metni text area'ya ve message box'a yaz
                    self.text_area.delete(1.0, tk.END)
                    self.text_area.insert(tk.END, transcribed_text)
                    messagebox.showinfo("Dönüştürülen Metin", transcribed_text)
                else:
                    messagebox.showwarning("Uyarı", "Dönüştürülen metin boş!")
            else:
                messagebox.showerror("Hata", "Transkripsiyon dosyası bulunamadı!")
            
            # UI'ı sıfırla
            self.record_button.configure(state='normal')
            self.status_label.configure(text="Hazır")
            
        except Exception as e:
            messagebox.showerror("Hata", f"Bir hata oluştu: {str(e)}")
            self.record_button.configure(state='normal')
            self.status_label.configure(text="Hata!")

    def save_to_word(self):
        """Metni Word dosyasına kaydetme"""
        try:
            filename = self.filename_entry.get()
            if not filename:
                messagebox.showerror("Hata", "Lütfen bir dosya adı girin!")
                return

            text = self.text_area.get(1.0, tk.END).strip()
            if not text:
                messagebox.showerror("Hata", "Dönüştürülmüş metin bulunamadı!")
                return

            # Word dosyası oluştur
            doc = Document()
            doc.add_heading('Ses Kaydı Dökümü', 0)
            doc.add_paragraph(f'Oluşturulma Tarihi: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
            doc.add_paragraph(text)

            # Dosyayı kaydet
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(self.output_dir, f"{filename}_{timestamp}.docx")
            doc.save(output_path)

            success_message = f"Dosya başarıyla kaydedildi!\nKonum: {os.path.abspath(output_path)}"
            messagebox.showinfo("Başarılı", success_message)
            
            # Dosya konumunu aç
            os.startfile(os.path.dirname(os.path.abspath(output_path)))

        except Exception as e:
            messagebox.showerror("Hata", f"Dosya kaydedilirken bir hata oluştu: {str(e)}")

    def cleanup(self):
        """Kaynakları temizle"""
        try:
            if hasattr(self, 'recorder'):
                self.recorder.stop_recording()
            
            # Geçici dosyaları sil
            if os.path.exists(self.transcript_path):
                os.remove(self.transcript_path)
        except Exception as e:
            print(f"Temizleme hatası: {str(e)}")

    def on_closing(self):
        """Uygulama kapatılırken temizlik"""
        if self.is_recording:
            if messagebox.askokcancel("Çıkış", "Kayıt devam ediyor. Çıkmak istediğinizden emin misiniz?"):
                self.cleanup()
                self.root.destroy()
        else:
            self.cleanup()
            self.root.destroy()

def main():
    root = tk.Tk()
    app = VoiceRecordingApp(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)
    root.mainloop()

if __name__ == "__main__":
    main()
